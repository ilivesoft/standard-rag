# 문서 파싱 모듈 - PDF, DOCX, TXT, MD, HTML, 이미지 지원
import re
from pathlib import Path
from pipeline.exceptions import UnsupportedFileTypeError, EmptyFileError

# 지원하는 파일 확장자 목록 (이미지 형식 포함)
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".html",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tiff",
}

# 이미지 확장자 집합
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


class DocumentParser:
    """다양한 문서 형식에서 텍스트를 추출하는 파서 클래스"""

    def __init__(self) -> None:
        # OCR 리더는 처음 사용 시 지연 초기화 (lazy loading)
        self._ocr_reader = None

    def _get_ocr_reader(self):
        """EasyOCR 리더를 지연 초기화하여 반환합니다.

        설정에서 OCR 언어 목록과 GPU 사용 여부를 읽어 초기화합니다.
        한 번 초기화된 이후에는 캐시된 인스턴스를 재사용합니다.
        """
        if self._ocr_reader is None:
            import easyocr
            from config.settings import settings

            self._ocr_reader = easyocr.Reader(
                settings.OCR_LANGUAGES,
                gpu=settings.OCR_GPU,
            )
        return self._ocr_reader

    def parse(self, file_path: Path) -> str:
        """파일 형식을 자동 감지하여 텍스트를 추출합니다.

        Args:
            file_path: 파싱할 파일 경로

        Returns:
            추출된 텍스트 문자열

        Raises:
            FileNotFoundError: 파일이 존재하지 않을 때
            UnsupportedFileTypeError: 지원하지 않는 파일 형식일 때
            EmptyFileError: 파일이 비어있거나 텍스트를 추출할 수 없을 때
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        if not self.is_supported(file_path):
            raise UnsupportedFileTypeError(
                f"지원하지 않는 파일 형식입니다: {file_path.suffix}"
            )

        ext = file_path.suffix.lower()

        if ext == ".pdf":
            text = self._parse_pdf(file_path)
        elif ext == ".docx":
            text = self._parse_docx(file_path)
        elif ext == ".html":
            text = self._parse_html(file_path)
        elif ext in _IMAGE_EXTENSIONS:
            # 이미지 파일은 OCR로 텍스트 추출
            text = self._parse_image(file_path)
        else:
            # .txt, .md
            text = self._parse_text(file_path)

        # 복잡한 문서 형식에서 빈 결과 시 UnstructuredFileLoader로 폴백
        _unstructured_fallback_exts = {".pdf", ".docx", ".html"}
        if (not text or not text.strip()) and ext in _unstructured_fallback_exts:
            try:
                text = self._parse_with_unstructured(file_path)
            except Exception:
                pass

        if not text or not text.strip():
            raise EmptyFileError(f"파일에서 텍스트를 추출할 수 없습니다: {file_path}")

        return text

    def _parse_pdf(self, file_path: Path) -> str:
        """PDF 파일에서 텍스트를 추출합니다.

        PyMuPDF로 네이티브 텍스트를 추출하고, 스캔 페이지는 OCR로 처리합니다.
        """
        import fitz  # PyMuPDF

        text_parts = []
        with fitz.open(str(file_path)) as doc:
            for page in doc:
                # 먼저 PDF 네이티브 텍스트 추출 시도
                page_text = page.get_text()

                if page_text and page_text.strip():
                    # 네이티브 텍스트가 있으면 그대로 사용
                    text_parts.append(page_text)
                else:
                    # 텍스트가 없는 페이지(스캔 PDF)는 이미지로 렌더링 후 OCR 수행
                    pix = page.get_pixmap()
                    img_bytes = pix.tobytes("png")
                    ocr_text = self._ocr_from_bytes(img_bytes)
                    if ocr_text:
                        text_parts.append(ocr_text)

        return "\n".join(text_parts)

    def _parse_docx(self, file_path: Path) -> str:
        """python-docx를 사용하여 DOCX 파일에서 텍스트를 추출합니다.

        단락과 테이블 내용을 모두 추출합니다. 병합 셀은 중복 제거 처리합니다.
        """
        import docx

        doc = docx.Document(str(file_path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # 병합 셀로 인한 연속 중복 제거
                deduped = [row_texts[0]] if row_texts else []
                for t in row_texts[1:]:
                    if t != deduped[-1]:
                        deduped.append(t)
                if deduped:
                    text_parts.append(" | ".join(deduped))

        return "\n".join(text_parts)

    def _parse_text(self, file_path: Path) -> str:
        """TXT 또는 MD 파일을 읽어 텍스트를 반환합니다."""
        return file_path.read_text(encoding="utf-8")

    def _parse_html(self, file_path: Path) -> str:
        """HTML 태그를 제거하고 순수 텍스트만 추출합니다."""
        html_content = file_path.read_text(encoding="utf-8")
        # HTML 태그 제거
        text = re.sub(r"<[^>]+>", " ", html_content)
        # 연속된 공백 정규화
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def _parse_image(self, file_path: Path) -> str:
        """EasyOCR을 사용하여 이미지 파일에서 텍스트를 추출합니다.

        Args:
            file_path: OCR 처리할 이미지 파일 경로

        Returns:
            인식된 텍스트를 줄바꿈으로 연결한 문자열
        """
        # Windows 한글 경로 문제: cv2.imread는 ANSI API로 열어 한글 경로를 처리 못함
        # 파일을 bytes로 읽은 뒤 numpy 배열로 변환하여 전달
        return self._ocr_from_bytes(file_path.read_bytes())

    def _ocr_from_bytes(self, img_bytes: bytes) -> str:
        """바이트 형태의 이미지 데이터에서 EasyOCR로 텍스트를 추출합니다.

        스캔된 PDF 페이지처럼 픽스맵(pixmap)에서 PNG 바이트로 변환된
        이미지에 대한 OCR 폴백 처리에 사용됩니다.

        Args:
            img_bytes: PNG 형식으로 인코딩된 이미지 바이트 데이터

        Returns:
            인식된 텍스트를 줄바꿈으로 연결한 문자열
        """
        import numpy as np
        import cv2

        reader = self._get_ocr_reader()
        # bytes → numpy 배열로 변환하여 cv2.imread 경로 문제 우회
        nparr = np.frombuffer(img_bytes, dtype=np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        results = reader.readtext(img, detail=0)
        return "\n".join(results)

    def _parse_with_unstructured(self, file_path: Path) -> str:
        """UnstructuredFileLoader를 사용하여 파일에서 텍스트를 추출합니다.

        langchain-community의 UnstructuredFileLoader를 활용하여
        다양한 파일 형식에서 텍스트를 추출합니다. 기본 파서의 폴백으로 사용됩니다.

        Args:
            file_path: 파싱할 파일 경로

        Returns:
            추출된 텍스트 문자열 (문서 단위로 개행 구분)
        """
        from langchain_community.document_loaders import UnstructuredFileLoader

        loader = UnstructuredFileLoader(str(file_path), mode="single")
        docs = loader.load()
        return "\n\n".join(doc.page_content for doc in docs if doc.page_content.strip())

    def is_supported(self, file_path: Path) -> bool:
        """파일 확장자가 지원되는 형식인지 확인합니다."""
        return Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS
